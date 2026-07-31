import os
from dotenv import load_dotenv
from google import genai
from docx import Document

load_dotenv()

client = genai.Client(
    api_key=os.getenv("GEMINI_API_KEY")
)

def procesar_audio_a_informe(ruta_audio):
    print("Subiendo audio...")

    audio_file = client.files.upload(file=ruta_audio)

    prompt = """
    Eres un asistente contable experto.

    Escucha atentamente este audio y redacta un INFORME CONTABLE profesional.

    El informe debe contener:

    1. Título.
    2. Resumen ejecutivo.
    3. Puntos clave.
    4. Tareas pendientes.

    Utiliza lenguaje formal y bien estructurado.
    """

    print("Generando informe...")

    response = client.models.generate_content(
    model="gemini-3.6-flash",
    contents=[audio_file, prompt]
)

    client.files.delete(name=audio_file.name)

    return response.text


def guardar_word(texto, nombre_archivo):
    doc = Document()
    doc.add_heading("Informe Contable", level=1)
    doc.add_paragraph(texto)
    doc.save(nombre_archivo)


if __name__ == "__main__":

    archivo_audio = input("Ingrese la ruta del archivo de audio: ").strip()

    if not os.path.exists(archivo_audio):
        print("❌ El archivo no existe.")
        exit()

    try:
        informe = procesar_audio_a_informe(archivo_audio)

        nombre_salida = input(
            "Nombre del archivo Word (sin .docx): "
        ).strip()

        if nombre_salida == "":
            nombre_salida = "informe_generado"

        guardar_word(informe, nombre_salida + ".docx")

        print(f"\n✅ Informe guardado como: {nombre_salida}.docx")

    except Exception as e:
        print("\nOcurrió un error:")
        print(e)